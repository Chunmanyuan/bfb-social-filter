from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT_DIR = Path(__file__).resolve().parents[2]
OUTPUT_PATH = ROOT_DIR / "社媒筛选器使用说明与注意事项_修正版.docx"


def set_run_font(run, name: str, size: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_heading(doc: Document, text: str, size: int = 13) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "Microsoft YaHei UI", size, bold=True)


def add_paragraph(doc: Document, text: str, size: int = 11, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "Microsoft YaHei UI", size, italic=italic)


def main() -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei UI"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei UI")
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title_run = title.add_run("社媒筛选器使用说明与注意事项")
    set_run_font(title_run, "Microsoft YaHei UI", 16, bold=True)

    add_paragraph(doc, "本说明面向首次接触本程序的普通用户，建议先完整阅读一次，再开始使用。")

    add_heading(doc, "一、文件说明")
    add_paragraph(doc, "1. 运行前安装向导.exe：首次使用时运行，用于检查 Chrome、检查 Python、安装或修复运行环境，并引导登录平台账号。")
    add_paragraph(doc, "2. 开始运行.exe：日常使用时运行，用于启动本地服务、打开筛选页面，并在右下角托盘常驻。")
    add_paragraph(doc, "3. 程序文件：程序核心目录，请勿随意删除、重命名，或单独移动其中内容。")
    add_paragraph(doc, "4. 爬虫与截图筛选程序：手动备份目录，日常使用无需打开。")

    add_heading(doc, "二、首次使用步骤")
    add_paragraph(doc, "5. 双击运行“运行前安装向导.exe”。")
    add_paragraph(doc, "6. 按提示完成环境检查与安装。首次安装时速度可能较慢，请耐心等待，不要中途关闭。")
    add_paragraph(doc, "7. 安装完成后，程序会提示打开小红书和 B 站登录窗口。")
    add_paragraph(doc, "8. 请在两个平台都完成登录。登录成功后，请保持页面打开至少 10 秒，再关闭窗口。")
    add_paragraph(doc, "9. 返回安装引导界面，点击“我已完成登录”。完成后即可进入日常使用阶段。")

    add_heading(doc, "三、日常使用步骤")
    add_paragraph(doc, "10. 双击运行“开始运行.exe”。")
    add_paragraph(doc, "11. 程序启动后会在电脑右下角托盘显示“社媒筛选器”图标，并自动打开浏览器页面。")
    add_paragraph(doc, "12. 如果浏览器页面被误关，可再次双击“开始运行.exe”重新打开页面。")
    add_paragraph(doc, "13. 如需完全退出，请在右下角托盘图标上右键，选择退出。")

    add_heading(doc, "四、注意事项")
    add_paragraph(doc, "14. 不要只单独移动两个 EXE 文件。正确做法是保持它们与“程序文件”文件夹处于同一层目录。")
    add_paragraph(doc, "15. 不要随意删除“程序文件”中的内容，否则程序可能无法启动。")
    add_paragraph(doc, "16. 首次安装、首次登录、首次启动时，速度会比平时慢一些，这属于正常现象。")
    add_paragraph(doc, "17. 如果安装引导提示缺少 Chrome 或 Python，请按界面提示处理后再重新运行。")
    add_paragraph(doc, "18. 如果页面打不开、数据加载失败，优先检查是否已经完成安装、是否已经登录平台账号，以及是否误退出了程序。")
    add_paragraph(doc, "19. 如遇异常，建议先关闭程序后重新运行；如仍无法解决，再联系维护人员。")

    add_heading(doc, "五、给用户的简单记忆方式")
    add_paragraph(doc, "20. 首次使用：先点“运行前安装向导.exe”。")
    add_paragraph(doc, "21. 以后使用：直接点“开始运行.exe”。")
    add_paragraph(doc, "22. 完全退出：到右下角托盘图标里退出。")

    add_paragraph(doc, "当前版本为初稿，如需用于正式分发，建议后续再补充常见问题、报错处理和联系方式。", italic=True)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
